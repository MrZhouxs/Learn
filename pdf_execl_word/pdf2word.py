# !/usr/bin/python3
# coding=utf-8
# @Time   : 2019/3/1/0001 14:45:40
# Author  : little star
# Func: pdf转换成word 默认只能转换中文
from io import StringIO

from docx import Document
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, process_pdf


def read_from_pdf(file_path):
    """
    获取pdf的内容，返回字符串
    :param file_path: pdf文件路径
    :return:
    """
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()

        content = return_str.getvalue()
        return_str.close()
        return content


def save_text_to_word(content, file_path):
    """
    将pdf的内容写入到word文档中
    :param content: pdf内容
    :param file_path: word文档存放的路径
    :return:
    """
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content):
    """
    去除控制符
    控制字符就是ASCII码在32以下的，所以我们使用str的translate方法，把32以下的字符移除
    :param content:
    :return:
    """
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)


def pdf_to_word(pdf_file_path, word_file_path):
    content = read_from_pdf(pdf_file_path)
    save_text_to_word(content, word_file_path)


if __name__ == '__main__':
    pdf_to_word("./test.pdf", "./test.docx")
